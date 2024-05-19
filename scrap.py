import requests
from bs4 import BeautifulSoup
from docx import Document

r= requests.get('https://www.javatpoint.com/python-interview-questions')
print(r)
soup = BeautifulSoup(r.content,'html.parser')
s= soup.find('div', class_ = 'onlycontent')
questions= s.find_all('h3')
answers= s.find_all('p')
doc= Document()
doc.add_heading('PYTHON INTERVIEW QUESTIONS', level=1)
for question in questions:
    doc.add_paragraph(question)

# first_question_encountered = False
# for question,answer in zip(questions,answers):
#     if not first_question_encountered:
#         first_question_encountered = True
#         continue
    
#     doc.add_heading(question.text.strip(), level=2)
#     doc.add_paragraph(answer.text.strip())
doc.save('python_interview_questions.docx')    